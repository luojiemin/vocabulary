import streamlit as st
st.set_page_config(page_title="英语词汇扩展系统", layout="centered")

import pandas as pd
import easyocr
from PIL import Image
import numpy as np
import io
from docx import Document

@st.cache_resource
def load_reader():
    return easyocr.Reader(['en', 'ch_sim'])

reader = load_reader()

st.title("📘 高三常忘英语词汇扩展学习系统")

uploaded_files = st.file_uploader("上传中英文截图（支持多张）", type=["png", "jpg", "jpeg"], accept_multiple_files=True)

results = []

def enrich_word_data(word, meaning):
    sample_dict = {
        "violate": {
            "pos": "v.",
            "collocation": "violate the law（违法）",
            "example": "He was fined for violating traffic rules.（他因违反交通规则被罚款。）",
            "derivatives": "violates, violating, violated, violation (n.), violator (n.)",
            "confusing": "violet（紫罗兰）, violent（暴力的）"
        }
    }
    data = sample_dict.get(word.lower(), {})
    return {
        "word": word,
        "pos": data.get("pos", ""),
        "cn_meaning": meaning,
        "collocation": data.get("collocation", ""),
        "example": data.get("example", ""),
        "derivatives": data.get("derivatives", ""),
        "confusing": data.get("confusing", "")
    }

def is_chinese(text):
    return any('\u4e00' <= ch <= '\u9fff' for ch in text)

if uploaded_files:
    with st.spinner("正在识别图像并提取词汇，请稍候..."):
        for file in uploaded_files:
            image = Image.open(file)
            result = reader.readtext(np.array(image))
            lines = [line[1] for line in result]
            st.text(f"DEBUG: 共识别到 {len(lines)} 行文字")
            st.write(lines)  # 调试显示所有文本

            for i in range(len(lines) - 1):
                if lines[i].isalpha() and is_chinese(lines[i + 1]):
                    word = lines[i]
                    meaning = lines[i + 1]
                    result_data = enrich_word_data(word, meaning)
                    results.append(result_data)

if not results and not uploaded_files:
    results = [enrich_word_data("violate", "违反")]

if results:
    df = pd.DataFrame(results)
    st.success(f"共提取到 {len(results)} 个词汇")
    st.dataframe(df)

    doc = Document()
    doc.add_heading("高三英语常忘词扩展记忆手册", level=1)
    for item in results:
        doc.add_paragraph(f"■ {item['word']}  {item['pos']}  {item['cn_meaning']}")
        doc.add_paragraph(f"常见搭配：{item['collocation']}")
        doc.add_paragraph(f"例句：{item['example']}")
        doc.add_paragraph(f"词形变化：{item['derivatives']}")
        doc.add_paragraph(f"形近词辨析：{item['confusing']}")
        doc.add_paragraph("")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button("📄 下载 Word 文档", buffer, file_name="词汇扩展结果.docx")
else:
    st.warning("未能识别到词汇，请确认截图中是否包含英文单词和释义。")
